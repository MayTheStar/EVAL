import os
import fitz  # PyMuPDF
import pdfplumber
from docx import Document
import docx2txt

def is_scanned_pdf(file_path):
    """
    Detects if a PDF is likely scanned (no extractable text).
    Returns True if scanned, False if digital text exists.
    """
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            if page.extract_text() and page.extract_text().strip():
                return False
    return True

def parse_pdf(file_path):
    if is_scanned_pdf(file_path):
        raise ValueError(
            "⚠️ This PDF appears to be scanned (image-based). "
            "Text extraction requires OCR."
        )

    text_content = ""
    tables_content = []
    images_saved = []

    # Extract text + tables
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text_content += page.extract_text() or ""
            tables = page.extract_tables()
            if tables:
                tables_content.extend(tables)

    # Extract images
    pdf_doc = fitz.open(file_path)
    for page_index in range(len(pdf_doc)):
        for img_index, img in enumerate(pdf_doc.get_page_images(page_index)):
            xref = img[0]
            pix = fitz.Pixmap(pdf_doc, xref)
            img_name = f"extracted_image_page{page_index+1}_{img_index+1}.png"
            pix.save(img_name)
            images_saved.append(img_name)

    return {"text": text_content, "tables": tables_content, "images": images_saved}

def parse_docx(file_path):
    doc = Document(file_path)
    text_content = "\n".join([p.text for p in doc.paragraphs])

    tables_content = []
    for table in doc.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        tables_content.append(rows)

    return {"text": text_content, "tables": tables_content, "images": []}

def parse_doc(file_path):
    text_content = docx2txt.process(file_path)
    return {"text": text_content, "tables": [], "images": []}

def parse_document(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    elif ext == ".doc":
        return parse_doc(file_path)
    else:
        raise ValueError("Unsupported file type. Use PDF, DOC, or DOCX.")

# Example usage
if __name__ == "__main__":
    file_path = "sample.pdf"  # replace with your file
    result = parse_document(file_path)
    print(result["text"][:1000], "\n...")  # preview text
    print("Tables:", result["tables"])
    print("Images saved:", result["images"])
