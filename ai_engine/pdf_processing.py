import os
import re
import json
from pathlib import Path
from typing import List, Dict, Any

import fitz  # PyMuPDF
import pdfplumber
import pytesseract
from pdf2image import convert_from_path
from docx import Document
import docx2txt

# ----------------------------
# Detect Scanned PDFs
# ----------------------------
def is_scanned_pdf(file_path: str) -> bool:
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text and text.strip():
                return False
    return True

# ----------------------------
# OCR for Scanned PDFs
# ----------------------------
def ocr_pdf(file_path: str) -> List[str]:
    pages_text = []
    images = convert_from_path(file_path)
    for img in images:
        pages_text.append(pytesseract.image_to_string(img))
    return pages_text

# ----------------------------
# Clean extracted text
# ----------------------------
def clean_text(text: str) -> str:
    text = text.replace("\n", " ")
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

# ----------------------------
# Extract Digital PDF
# ----------------------------
def extract_pdf(file_path: str) -> List[Dict[str, Any]]:
    pages_data = []

    if is_scanned_pdf(file_path):
        ocr_texts = ocr_pdf(file_path)
        for idx, txt in enumerate(ocr_texts):
            pages_data.append({"page_number": idx + 1, "text": txt, "tables": []})
        return pages_data

    with pdfplumber.open(file_path) as pdf:
        for idx, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            tables = page.extract_tables()

            pages_data.append({
                "page_number": idx + 1,
                "text": text,
                "tables": [
                    {
                        "table_index": t + 1,
                        "rows": len(tbl),
                        "columns": len(tbl[0]) if tbl else 0,
                        "data": tbl
                    } for t, tbl in enumerate(tables)
                ]
            })
    return pages_data

# ----------------------------
# Extract DOCX / DOC
# ----------------------------
def extract_docx(file_path: str):
    doc = Document(file_path)
    tables = []
    for i, table in enumerate(doc.tables):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        tables.append({"table_index": i + 1, "rows": len(rows), "columns": len(rows[0]) if rows else 0, "data": rows})
    return [{"page_number": 1, "text": "\n".join(p.text for p in doc.paragraphs), "tables": tables}]

def extract_doc(file_path: str):
    return [{"page_number": 1, "text": docx2txt.process(file_path), "tables": []}]

# ----------------------------
# Full Pipeline
# ----------------------------
def parse_document(file_path: str):

    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        pages = extract_pdf(file_path)
    elif ext == ".docx":
        pages = extract_docx(file_path)
    elif ext == ".doc":
        pages = extract_doc(file_path)
    else:
        raise ValueError("Unsupported file format!")

    for page in pages:
        page["cleaned_text"] = clean_text(page["text"])

    return pages